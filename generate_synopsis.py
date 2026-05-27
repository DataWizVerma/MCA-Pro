"""
generate_synopsis.py
--------------------
Generates a concise Synopsis (.docx) for Kumar Verma's MCA Final Year
Major Project - Chandigarh University (23ONMCR-753).

As per university guidelines, synopsis contains ONLY:
  1. Title of the Project
  2. Objective
  3. Resource Required  (Hardware + Software)

Run:    python generate_synopsis.py
Output: Synopsis_Kumar_Verma_MCA.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────── HELPER FUNCTIONS ────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set background colour of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def center_para(doc, text, size, bold=False, color="000000", space_before=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(
        int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
    return p


def section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold      = True
    run.underline = True
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(31, 56, 100)


def body_text(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'


def bullet_point(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'


def make_table(doc, headers, rows, col_widths=(2.5, 4.0)):
    table = doc.add_table(rows=1, cols=2)
    table.style     = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].width = Inches(col_widths[i])
        set_cell_bg(hdr[i], '1F3864')
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11.5)
        run.font.name = 'Times New Roman'
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for idx, (c1, c2) in enumerate(rows):
        cells = table.add_row().cells
        cells[0].width = Inches(col_widths[0])
        cells[1].width = Inches(col_widths[1])
        bg = 'EBF5FB' if idx % 2 == 0 else 'FFFFFF'
        set_cell_bg(cells[0], bg)
        set_cell_bg(cells[1], bg)
        for cell, val in zip(cells, [c1, c2]):
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(11.5)
            run.font.name = 'Times New Roman'
    doc.add_paragraph()


# ─────────────────────────── DOCUMENT SETUP ──────────────────────────────────

doc = Document()

# Margins: 1.25 inches as per university guideline
for sec in doc.sections:
    sec.top_margin    = Inches(1.25)
    sec.bottom_margin = Inches(1.25)
    sec.left_margin   = Inches(1.25)
    sec.right_margin  = Inches(1.25)

# Default style
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = Pt(18)   # 1.5 line spacing


# ─────────────────────────── TITLE PAGE ──────────────────────────────────────

doc.add_paragraph('\n\n')

center_para(doc, "CHANDIGARH UNIVERSITY",
            size=16, bold=True, color="1F3864")
center_para(doc, "Centre for Distance and Online Education",
            size=12, color="1F3864")
center_para(doc, "Master of Computer Applications (MCA) - Fourth Semester",
            size=12, color="555555")

doc.add_paragraph()
center_para(doc, "SYNOPSIS", size=22, bold=True, color="C0392B")
center_para(doc, "of Major Project  |  Subject Code: 23ONMCR-753",
            size=11, color="555555")
doc.add_paragraph()

# Decorative line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
for side in ('top', 'bottom'):
    el = OxmlElement(f'w:{side}')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), '8')
    el.set(qn('w:space'), '4')
    el.set(qn('w:color'), '1F3864')
    pBdr.append(el)
pPr.append(pBdr)
doc.add_paragraph()

# Project title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Digital Image Forensics System\nfor Tampering Detection")
r.bold = True
r.font.size = Pt(15)
r.font.name = 'Times New Roman'
r.font.color.rgb = RGBColor(31, 56, 100)

center_para(doc, "Using Error Level Analysis and Metadata Analysis",
            size=12, color="555555")

doc.add_paragraph('\n\n')

# Student info table
info_table = doc.add_table(rows=5, cols=2)
info_table.style     = 'Table Grid'
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, value) in enumerate([
    ("Submitted By",  "Kumar Verma"),
    ("Programme",     "Master of Computer Applications (MCA)"),
    ("Semester",      "Fourth Semester"),
    ("Subject Code",  "23ONMCR-753"),
    ("Session",       "2024 - 25"),
]):
    row = info_table.rows[i]
    row.cells[0].width = Inches(2.0)
    row.cells[1].width = Inches(4.5)
    set_cell_bg(row.cells[0], '1F3864')
    r1 = row.cells[0].paragraphs[0].add_run(label)
    r1.bold = True; r1.font.size = Pt(11); r1.font.name = 'Times New Roman'
    r1.font.color.rgb = RGBColor(255, 255, 255)
    r2 = row.cells[1].paragraphs[0].add_run(value)
    r2.font.size = Pt(11); r2.font.name = 'Times New Roman'

doc.add_page_break()


# ─────────────────────────── SECTION 1: TITLE ────────────────────────────────

section_heading(doc, "1.  Title of the Project")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "Digital Image Forensics System for Tampering Detection\n"
    "Using Error Level Analysis and Metadata Analysis"
)
r.bold = True
r.font.size = Pt(13)
r.font.name = 'Times New Roman'
r.font.color.rgb = RGBColor(31, 56, 100)
doc.add_paragraph()


# ─────────────────────────── SECTION 2: OBJECTIVE ────────────────────────────

section_heading(doc, "2.  Objective")

body_text(doc,
    "The objective of this project is to design and develop a web-based Digital Image Forensics "
    "System that automatically detects whether a digital photograph has been tampered or manipulated "
    "using Artificial Intelligence and multiple forensic techniques. The specific objectives are:"
)

objectives = [
    # Fix 1: No accuracy percentages — avoids viva questions on training details
    "To detect image tampering using Error Level Analysis (ELA) with a pre-trained DenseNet121 "
    "deep learning model, and to visualise suspicious regions through ELA heatmaps.",

    "To extract and analyse EXIF metadata embedded in photographs, including camera information, "
    "GPS coordinates, capture timestamp, and post-processing software details.",

    "To validate image authenticity by cross-referencing EXIF location and timestamp data against "
    "historical weather records and independently classifying visible weather using a CNN model.",

    "To perform additional forensic checks including cryptographic hash computation (MD5 / SHA-256), "
    "Copy-Move forgery detection, JPEG double-compression analysis, and noise consistency analysis — "
    "and to present all results in a unified web dashboard with a downloadable forensic report.",
]

for obj in objectives:
    bullet_point(doc, obj)


# ─────────────────────── SECTION 3: RESOURCE REQUIRED ────────────────────────

section_heading(doc, "3.  Resource Required")

# Sub-heading A: Hardware
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after  = Pt(4)
r = p.add_run("A.  Hardware Requirements")
r.bold = True; r.font.size = Pt(12); r.font.name = 'Times New Roman'

make_table(doc,
    ["Component", "Specification"],
    [
        ("Processor",    "Intel Core i3 (4th Generation) or above"),
        ("RAM",          "Minimum 4 GB  (8 GB Recommended)"),
        ("Hard Disk",    "Minimum 5 GB Free Space"),
        ("Display",      "1024 x 768 Resolution or higher"),
        ("Internet",     "Required for Weather API and Geocoding"),
        ("Graphics Card","Optional — CPU inference is supported"),
    ],
    col_widths=(2.2, 4.3)
)

# Sub-heading B: Software
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(4)
r = p.add_run("B.  Software Requirements")
r.bold = True; r.font.size = Pt(12); r.font.name = 'Times New Roman'

make_table(doc,
    ["Category", "Tool / Technology"],
    [
        ("Operating System",         "Windows 10 / 11  or  Linux Ubuntu 20.04+"),
        ("Front-End Tool",           "Streamlit (Python Web Framework)"),
        ("Back-End / Storage",       "Local File System + HDF5 (.h5) Model Files"),
        ("Coding Language",          "Python 3.10 or 3.11"),    # Fix 2: stable TF support
        ("IDE",                      "Visual Studio Code / PyCharm"),
        ("Deep Learning Framework",  "TensorFlow 2.x, tf_keras"),
        ("Image Processing",         "OpenCV, Pillow (PIL), piexif"),
        ("Data Analysis",            "NumPy, Pandas, Matplotlib"),
        ("Geocoding Library",        "geopy  (Nominatim — OpenStreetMap)"),
        ("Weather Data API",         "Open-Meteo ERA5 Historical API  (Free)"),
        ("Internet Technology",      "HTTP / HTTPS  (REST API),  TCP/IP"),
        ("Training Dataset",         "CASIA 2.0  Image Forgery Benchmark"),
    ],
    col_widths=(2.2, 4.3)
)


# ─────────────────────────── SIGNATURE BLOCK ─────────────────────────────────
# Fix 3: Proper guide block with Name / Designation / Signature lines

doc.add_paragraph()
sig = doc.add_table(rows=2, cols=2)
sig.style     = 'Table Grid'
sig.alignment = WD_TABLE_ALIGNMENT.CENTER

sig_data = [
    # Header row
    ("Submitted By", "Approved By"),
    # Content row — student left, guide right with proper fields
    (
        "Kumar Verma\n"
        "MCA  |  Fourth Semester\n"
        "Chandigarh University\n"
        "Session: 2024 - 25\n\n"
        "Signature: _________________",

        "Name: ______________________\n"
        "Designation: ________________\n"
        "Department: _________________\n\n"
        "Signature: _________________"
    ),
]

for r_idx, (left, right) in enumerate(sig_data):
    for c_idx, val in enumerate([left, right]):
        cell = sig.cell(r_idx, c_idx)
        cell.paragraphs[0].clear()
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if r_idx == 0:
            set_cell_bg(cell, '1F3864')
            run = cell.paragraphs[0].add_run(val)
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        else:
            run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'


# ─────────────────────────── SAVE FILE ───────────────────────────────────────

out = "Synopsis_Kumar_Verma_MCA.docx"
doc.save(out)
print("Saved: " + out)
