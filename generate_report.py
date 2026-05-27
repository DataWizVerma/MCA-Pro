"""
generate_report.py
------------------
Generates the complete MCA Final Year Project Report (.docx) for
Kumar Verma - Chandigarh University (23ONMCR-753).

Sections as per university guidelines:
  A. Title Page        E. Abstract          I. Design
  B. Certificate       F. Table of Contents J. Coding & Implementation
  C. Declaration       G. Introduction      K. Testing
  D. Acknowledgement   H. SDLC              L. Application
                                            M. Conclusion
                                            N. Bibliography (APA)

Run:    python generate_report.py
Output: Project_Report_Kumar_Verma_MCA.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def add_page_break(doc):
    doc.add_page_break()


def centered(doc, text, size=12, bold=False, color="000000",
             space_before=6, space_after=6, italic=False):
    p   = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(
        int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
    return p


def heading1(doc, text):
    """Chapter heading - large, dark blue, centered."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(12)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(31, 56, 100)
    # Bottom border
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '8')
    bot.set(qn('w:space'), '4')
    bot.set(qn('w:color'), '1F3864')
    pBdr.append(bot)
    pPr.append(pBdr)


def heading2(doc, text):
    """Section heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(31, 56, 100)


def heading3(doc, text):
    """Sub-section heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.underline = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)


def body(doc, text, justify=True):
    """Normal body paragraph, 1.5 line spacing."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before  = Pt(3)
    p.paragraph_format.space_after   = Pt(6)
    p.paragraph_format.line_spacing  = Pt(22)   # ~1.5 for 12pt
    if justify:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    return p


def bullet(doc, text, level=0):
    """Bullet point."""
    style = 'List Bullet'
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.25)
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'


def numbered_item(doc, number, text):
    """Numbered list item."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.35)
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run(f"{number}.  {text}")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'


def code_block(doc, text):
    """Monospace code block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.line_spacing  = Pt(16)
    # Light grey background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F0F0F0')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)


def make_table(doc, headers, rows, col_widths=None, hdr_color='1F3864'):
    n  = len(headers)
    if col_widths is None:
        col_widths = [6.5 / n] * n
    t = doc.add_table(rows=1, cols=n)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].width = Inches(col_widths[i])
        set_cell_bg(hdr_cells[i], hdr_color)
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for idx, row in enumerate(rows):
        cells = t.add_row().cells
        bg = 'EBF5FB' if idx % 2 == 0 else 'FFFFFF'
        for i, val in enumerate(row):
            cells[i].width = Inches(col_widths[i])
            set_cell_bg(cells[i], bg)
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
    doc.add_paragraph()
    return t


def caption(doc, text):
    """Figure/table caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    run = p.add_run(text)
    run.italic    = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(80, 80, 80)


def separator(doc):
    """Thin horizontal rule."""
    p    = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bot)
    pPr.append(pBdr)


# =============================================================================
# DOCUMENT SETUP
# =============================================================================

doc = Document()

for sec in doc.sections:
    sec.top_margin    = Inches(1.25)
    sec.bottom_margin = Inches(1.25)
    sec.left_margin   = Inches(1.25)
    sec.right_margin  = Inches(1.25)

normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = Pt(22)


# =============================================================================
# A. TITLE PAGE
# =============================================================================

for _ in range(2):
    doc.add_paragraph()

centered(doc, "CHANDIGARH UNIVERSITY", 17, bold=True, color="1F3864", space_before=0)
centered(doc, "Centre for Distance and Online Education", 13, color="1F3864")
centered(doc, "Mohali, Punjab - 140413", 12, color="555555")

doc.add_paragraph()

# Decorative double line
for _ in range(2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ('top',):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'double')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), '1F3864')
        pBdr.append(el)
    pPr.append(pBdr)

doc.add_paragraph()
centered(doc, "PROJECT REPORT", 20, bold=True, color="C0392B")
centered(doc, "Submitted in Partial Fulfillment of the Requirements for the Degree of", 12, color="555555")
centered(doc, "Master of Computer Applications (MCA)", 14, bold=True, color="1F3864")
centered(doc, "Fourth Semester  |  Subject Code: 23ONMCR-753", 12, color="555555")

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("DIGITAL IMAGE FORENSICS SYSTEM\nFOR TAMPERING DETECTION")
r.bold = True; r.font.size = Pt(16); r.font.name = 'Times New Roman'
r.font.color.rgb = RGBColor(31, 56, 100)

centered(doc, "Using Error Level Analysis and Metadata Analysis", 12, italic=True, color="555555")

for _ in range(3):
    doc.add_paragraph()

# Info grid
t = doc.add_table(rows=6, cols=2)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (lbl, val) in enumerate([
    ("Submitted By",    "Kumar Verma"),
    ("Programme",       "Master of Computer Applications (MCA)"),
    ("Semester",        "Fourth Semester"),
    ("Subject Code",    "23ONMCR-753"),
    ("University",      "Chandigarh University"),
    ("Session",         "2024 - 25"),
]):
    row = t.rows[i]
    row.cells[0].width = Inches(2.0)
    row.cells[1].width = Inches(4.5)
    set_cell_bg(row.cells[0], '1F3864')
    r1 = row.cells[0].paragraphs[0].add_run(lbl)
    r1.bold = True; r1.font.size = Pt(11); r1.font.name = 'Times New Roman'
    r1.font.color.rgb = RGBColor(255, 255, 255)
    r2 = row.cells[1].paragraphs[0].add_run(val)
    r2.font.size = Pt(11); r2.font.name = 'Times New Roman'

add_page_break(doc)


# =============================================================================
# B. CERTIFICATE
# =============================================================================

heading1(doc, "B.  Certificate")

doc.add_paragraph()
body(doc,
    "This is to certify that the project report entitled \"Digital Image Forensics System for "
    "Tampering Detection Using Error Level Analysis and Metadata Analysis\" is a bonafide record "
    "of the project work carried out by Mr. Kumar Verma (MCA, Fourth Semester) at Chandigarh "
    "University, Centre for Distance and Online Education, in partial fulfillment of the requirements "
    "for the course Major Project (23ONMCR-753) during the academic session 2024-25."
)
body(doc,
    "This project has been carried out individually and has not been submitted to any other "
    "university or institution for the award of any degree or diploma."
)

for _ in range(3):
    doc.add_paragraph()

# Certificate signature table
ct = doc.add_table(rows=2, cols=2)
ct.style = 'Table Grid'
ct.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri, (l, r_) in enumerate([
    ("Project Guide",             "Examiner"),
    ("Name:   ______________________\n"
     "Designation: ___________________\n"
     "Department:  ___________________\n\n"
     "Signature:   ___________________\n"
     "Date:        ___________________",
     "Name:   ______________________\n"
     "Institution: ___________________\n\n\n"
     "Signature:   ___________________\n"
     "Date:        ___________________"),
]):
    for ci, val in enumerate([l, r_]):
        cell = ct.cell(ri, ci)
        cell.paragraphs[0].clear()
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if ri == 0:
            set_cell_bg(cell, '1F3864')
            run = cell.paragraphs[0].add_run(val)
            run.bold = True; run.font.color.rgb = RGBColor(255, 255, 255)
        else:
            run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(11); run.font.name = 'Times New Roman'

add_page_break(doc)


# =============================================================================
# C. DECLARATION
# =============================================================================

heading1(doc, "C.  Declaration")

doc.add_paragraph()
body(doc,
    "I, Kumar Verma, student of Master of Computer Applications (MCA), Fourth Semester, "
    "Chandigarh University, Centre for Distance and Online Education, hereby declare that "
    "the project report entitled \"Digital Image Forensics System for Tampering Detection "
    "Using Error Level Analysis and Metadata Analysis\" submitted for the Major Project "
    "(Subject Code: 23ONMCR-753) is my own original work."
)
body(doc,
    "I further declare that:"
)
numbered_item(doc, "1",
    "The project has been carried out individually under the guidance of my project supervisor.")
numbered_item(doc, "2",
    "The work presented in this report has not been submitted for any other degree, diploma, "
    "or certificate at this or any other university or institution.")
numbered_item(doc, "3",
    "All the information, ideas, and materials obtained from other sources have been duly "
    "acknowledged in the bibliography section of this report.")
numbered_item(doc, "4",
    "To the best of my knowledge, the project does not contain any material previously "
    "published or written by another person, except where due reference is given.")

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run("Name:      Kumar Verma\n"
                "Programme: Master of Computer Applications (MCA)\n"
                "Semester:  Fourth Semester\n"
                "Session:   2024-25\n\n"
                "Signature: _______________________\n"
                "Date:      _______________________")
run.font.size = Pt(12); run.font.name = 'Times New Roman'

add_page_break(doc)


# =============================================================================
# D. ACKNOWLEDGEMENT
# =============================================================================

heading1(doc, "D.  Acknowledgement")

doc.add_paragraph()
body(doc,
    "First and foremost, I express my sincere gratitude to Chandigarh University and the "
    "Centre for Distance and Online Education for providing me with the opportunity to "
    "undertake this major project as part of my MCA programme."
)
body(doc,
    "I am deeply thankful to my Project Guide for their continuous guidance, valuable "
    "suggestions, and encouragement throughout the development of this project. Their "
    "technical insights and constructive feedback were instrumental in shaping the direction "
    "and quality of this work."
)
body(doc,
    "I would also like to acknowledge the open-source community whose tools and libraries -- "
    "including Python, TensorFlow, OpenCV, and Streamlit -- formed the backbone of this "
    "project's implementation. The availability of the CASIA 2.0 dataset and the Open-Meteo "
    "ERA5 weather API made the forensic validation components possible."
)
body(doc,
    "A special word of thanks goes to my family for their constant moral support and "
    "encouragement. Their belief in my abilities gave me the strength to complete this "
    "project successfully."
)
body(doc,
    "Finally, I am grateful to all those who directly or indirectly contributed to the "
    "completion of this work.")

for _ in range(3):
    doc.add_paragraph()
centered(doc, "Kumar Verma", 12, bold=True)
centered(doc, "MCA -- Fourth Semester", 12)
centered(doc, "Chandigarh University", 12)
centered(doc, "Session: 2024-25", 12)

add_page_break(doc)


# =============================================================================
# E. ABSTRACT
# =============================================================================

heading1(doc, "E.  Abstract")

doc.add_paragraph()
body(doc,
    "Digital image tampering has emerged as a significant threat in the modern information age. "
    "With the widespread availability of sophisticated image editing software, it has become "
    "alarmingly easy to alter photographs without leaving any visible trace detectable by the "
    "human eye. This poses serious challenges across multiple domains including journalism, "
    "law enforcement, insurance, and social media -- where the authenticity of visual content "
    "is of paramount importance."
)
body(doc,
    "This project presents the design and implementation of a comprehensive web-based "
    "Digital Image Forensics System for Tampering Detection. The system integrates five "
    "independent forensic analysis modules into a unified platform, enabling multi-modal "
    "examination of JPEG digital photographs."
)
body(doc,
    "The primary analysis technique is Error Level Analysis (ELA), which exploits the "
    "differential compression history of edited versus unedited image regions. The ELA "
    "output is classified by a pre-trained DenseNet121 deep learning model trained on the "
    "CASIA 2.0 image forgery benchmark dataset. A complementary Weather Validation module "
    "cross-references GPS coordinates and capture timestamps extracted from EXIF metadata "
    "against historical weather records obtained via the Open-Meteo ERA5 API, and compares "
    "these with the weather conditions independently identified by a Convolutional Neural "
    "Network (CNN) weather classifier."
)
body(doc,
    "Additional forensic modules include cryptographic hash computation (MD5 and SHA-256) "
    "for file integrity verification, a block-based Copy-Move forgery detection algorithm, "
    "and JPEG quantization table analysis for detecting double compression -- a common "
    "artefact of post-capture editing. Noise consistency analysis across a 3x3 image grid "
    "further aids in detecting image splicing."
)
body(doc,
    "The system is deployed as a Streamlit web application featuring a dark-themed, "
    "multi-tab dashboard. Results are presented in both technical and plain-language formats, "
    "making the system accessible to non-technical users. A downloadable forensic report is "
    "also generated for documentation purposes."
)
body(doc,
    "The proposed system demonstrates that combining multiple independent forensic signals "
    "into a unified verdict significantly reduces the false positive rate of individual "
    "methods, providing a more reliable and robust authenticity assessment."
)

for _ in range(2):
    doc.add_paragraph()
make_table(doc,
    ["Keywords"],
    [("Digital Image Forensics, Error Level Analysis, ELA, EXIF Metadata, Image Tampering "
      "Detection, DenseNet121, Copy-Move Detection, Deep Learning, Streamlit, CASIA 2.0",)],
    col_widths=[6.5], hdr_color='2E4057'
)

add_page_break(doc)


# =============================================================================
# F. TABLE OF CONTENTS
# =============================================================================

heading1(doc, "F.  Table of Contents")

toc_items = [
    ("A.", "Title Page",                              "1"),
    ("B.", "Certificate",                             "2"),
    ("C.", "Declaration",                             "3"),
    ("D.", "Acknowledgement",                         "4"),
    ("E.", "Abstract",                                "5"),
    ("F.", "Table of Contents",                       "6"),
    ("G.", "Introduction",                            "7"),
    ("",   "  1.1  Background and Overview",          "7"),
    ("",   "  1.2  Problem Statement",                "9"),
    ("",   "  1.3  Objectives of the Project",        "10"),
    ("",   "  1.4  Scope of the Project",             "11"),
    ("",   "  1.5  Organisation of the Report",       "12"),
    ("H.", "SDLC of the Project",                     "13"),
    ("",   "  2.1  SDLC Overview",                    "13"),
    ("",   "  2.2  Waterfall Model",                  "14"),
    ("",   "  2.3  Phase-wise Description",           "15"),
    ("I.", "Design",                                  "19"),
    ("",   "  3.1  System Architecture",              "19"),
    ("",   "  3.2  Module Descriptions",              "20"),
    ("",   "  3.3  Data Flow Diagram",                "22"),
    ("",   "  3.4  Flowchart",                        "24"),
    ("",   "  3.5  Input / Output Design",            "25"),
    ("J.", "Coding and Implementation",               "27"),
    ("",   "  4.1  Development Environment",          "27"),
    ("",   "  4.2  Module 1 -- ELA Analysis",         "28"),
    ("",   "  4.3  Module 2 -- EXIF Metadata",        "32"),
    ("",   "  4.4  Module 3 -- Weather Validation",   "34"),
    ("",   "  4.5  Module 4 -- Cryptographic Hash",   "36"),
    ("",   "  4.6  Module 5 -- Copy-Move Detection",  "37"),
    ("",   "  4.7  Web Application Interface",        "39"),
    ("K.", "Testing",                                 "43"),
    ("",   "  5.1  Testing Strategy",                 "43"),
    ("",   "  5.2  Unit Testing",                     "44"),
    ("",   "  5.3  Integration Testing",              "46"),
    ("",   "  5.4  System Testing",                   "47"),
    ("",   "  5.5  Test Case Table",                  "48"),
    ("L.", "Application",                             "52"),
    ("",   "  6.1  Real-World Use Cases",             "52"),
    ("",   "  6.2  User Guide",                       "53"),
    ("M.", "Conclusion",                              "55"),
    ("",   "  7.1  Summary",                          "55"),
    ("",   "  7.2  Limitations",                      "56"),
    ("",   "  7.3  Future Scope",                     "57"),
    ("N.", "Bibliography (APA Style)",                "58"),
]

for sec, title, pg in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    is_main = sec != ""
    label = f"{sec}  {title}" if sec else title
    run = p.add_run(label)
    run.font.size = Pt(12 if is_main else 11)
    run.font.name = 'Times New Roman'
    run.bold = is_main
    if is_main:
        run.font.color.rgb = RGBColor(31, 56, 100)
    # right-aligned page number via tab
    p.add_run(f"\t{pg}").font.name = 'Times New Roman'

add_page_break(doc)


# =============================================================================
# G. INTRODUCTION
# =============================================================================

heading1(doc, "G.  Introduction")

heading2(doc, "1.1  Background and Overview")
body(doc,
    "Photography has undergone a transformative evolution since the invention of the daguerreotype "
    "in the nineteenth century. The transition from film-based to digital photography, and "
    "subsequently to smartphone photography, has democratised image capture to an unprecedented "
    "degree. Today, billions of photographs are taken and shared every day across social media "
    "platforms, news outlets, messaging applications, and professional databases."
)
body(doc,
    "While the proliferation of digital photography has brought immense social and economic "
    "benefits, it has simultaneously created a significant challenge: the ease with which "
    "digital images can be manipulated. Unlike their film-based predecessors -- where physical "
    "tampering left detectable traces -- digital photographs can be altered with a level of "
    "precision and invisibility that is undetectable by the unaided human eye."
)
body(doc,
    "Software tools such as Adobe Photoshop, GIMP (GNU Image Manipulation Program), "
    "and AI-powered image generation and editing platforms have become widely accessible. "
    "A skilled editor can insert, remove, clone, or composite objects within an image in "
    "minutes, leaving no visually apparent evidence of manipulation."
)
body(doc,
    "The consequences of undetected image forgery span multiple critical domains:"
)
bullet(doc, "Journalism and Media: Fake photographs are used to fabricate news stories and manipulate "
       "public opinion. A single viral tampered image can trigger social unrest or spread dangerous misinformation.")
bullet(doc, "Law Enforcement and Legal Proceedings: Photographic evidence is commonly submitted in courts "
       "of law. A tampered evidence photograph could lead to wrongful convictions or acquittals.")
bullet(doc, "Insurance Industry: Fraudulent insurance claims are frequently supported by manipulated "
       "photographs of accidents, property damage, or injuries, leading to significant financial losses.")
bullet(doc, "Academic and Scientific Research: Manipulated images in research papers -- particularly in "
       "life sciences -- constitute a form of academic fraud that can mislead entire research communities.")
bullet(doc, "Social Media and Personal Reputation: Individuals can be placed in fabricated contexts "
       "through image manipulation, leading to reputational damage, cyberbullying, and harassment.")

body(doc,
    "Digital image forensics is the scientific discipline that deals with the analysis of digital "
    "photographs to determine their authenticity and origin. It draws on principles from signal "
    "processing, information theory, machine learning, and cryptography to detect evidence of "
    "manipulation that is invisible to the human eye but detectable through computational means."
)
body(doc,
    "This project presents a comprehensive, multi-technique digital image forensics system "
    "implemented as an accessible web application. By combining five independent forensic "
    "techniques -- Error Level Analysis, EXIF Metadata Validation, Weather Cross-Referencing, "
    "Cryptographic Integrity Verification, and Copy-Move Detection -- the system provides a "
    "robust and accessible tool for assessing the authenticity of digital photographs."
)

heading2(doc, "1.2  Problem Statement")
body(doc,
    "Despite the growing prevalence and impact of digital image manipulation, there is a "
    "significant gap in the availability of free, accessible, and comprehensive image "
    "forensics tools for non-specialist users. The specific problems that motivate this "
    "project are as follows:"
)
numbered_item(doc, "1",
    "Image editing tools are widely available and easy to use, while tools for detecting "
    "manipulation are primarily limited to expensive commercial forensics software or "
    "command-line academic tools that require deep technical expertise.")
numbered_item(doc, "2",
    "Most existing forensic tools focus on a single detection technique (typically ELA "
    "alone), which is susceptible to false positives. A multi-technique approach that "
    "corroborates findings across independent methods is more reliable but is not "
    "commonly available in a unified platform.")
numbered_item(doc, "3",
    "JPEG compression artefacts introduced during image editing carry detectable forensic "
    "signatures, but analysing these requires computational image processing that most "
    "users cannot perform manually.")
numbered_item(doc, "4",
    "The EXIF metadata embedded in every photograph by cameras and smartphones contains "
    "rich forensic information -- including GPS coordinates and timestamps -- that can be "
    "cross-validated against external databases. However, this validation is rarely "
    "automated or integrated into existing tools.")
numbered_item(doc, "5",
    "Presenting forensic results in a format comprehensible to non-technical users -- "
    "such as journalists, legal professionals, and insurance investigators -- remains an "
    "unaddressed design challenge in existing forensics software.")

heading2(doc, "1.3  Objectives of the Project")
body(doc, "The following are the principal objectives of this project:")
numbered_item(doc, "1",
    "To detect image tampering using Error Level Analysis (ELA) with a pre-trained "
    "DenseNet121 deep learning model, and to visualise suspicious regions through "
    "ELA heatmap overlays.")
numbered_item(doc, "2",
    "To extract and analyse EXIF metadata embedded in photographs, including camera "
    "make and model, GPS coordinates, capture timestamp, and post-processing software "
    "details.")
numbered_item(doc, "3",
    "To validate image authenticity by cross-referencing EXIF location and timestamp "
    "data against historical weather records from the Open-Meteo ERA5 API and "
    "independently classifying visible weather conditions using a CNN model.")
numbered_item(doc, "4",
    "To perform comprehensive additional forensic checks -- including cryptographic hash "
    "computation, Copy-Move forgery detection, JPEG double-compression analysis, and "
    "noise consistency analysis -- and to present all results in a unified web dashboard "
    "with a downloadable forensic report.")

heading2(doc, "1.4  Scope of the Project")
body(doc,
    "The proposed system is designed to operate on JPEG (Joint Photographic Experts Group) "
    "format digital photographs. JPEG is the dominant image format in digital photography "
    "and is the format most susceptible to the compression-based forensic analysis techniques "
    "implemented in this system."
)
body(doc, "The system is applicable in the following contexts:")
bullet(doc, "Academic research and education in the field of digital media forensics")
bullet(doc, "Preliminary forensic investigation support for journalists and fact-checkers")
bullet(doc, "Supporting evidence evaluation in legal and insurance contexts")
bullet(doc, "Training and awareness programmes on digital media literacy")
body(doc,
    "The system is explicitly designed for academic and educational purposes. Its results "
    "are intended as supplementary forensic indicators and should not be used as the sole "
    "basis for legal decisions. Human expert review is recommended for all critical applications."
)

heading2(doc, "1.5  Organisation of the Report")
body(doc, "The remainder of this report is organised as follows:")
bullet(doc, "Chapter H -- SDLC: Describes the Software Development Life Cycle model adopted and the "
       "development phases of the project.")
bullet(doc, "Chapter I -- Design: Presents the system architecture, module design, data flow diagrams, "
       "and input/output specifications.")
bullet(doc, "Chapter J -- Coding and Implementation: Details the implementation of each module with "
       "code excerpts and explanations.")
bullet(doc, "Chapter K -- Testing: Documents the testing strategy, test cases, and results.")
bullet(doc, "Chapter L -- Application: Describes the real-world use cases and provides a user guide.")
bullet(doc, "Chapter M -- Conclusion: Summarises the project, discusses limitations, and outlines "
       "directions for future work.")
bullet(doc, "Chapter N -- Bibliography: Lists all references cited in APA format.")

add_page_break(doc)


# =============================================================================
# H. SDLC
# =============================================================================

heading1(doc, "H.  SDLC of the Project")

heading2(doc, "2.1  SDLC Overview")
body(doc,
    "The Software Development Life Cycle (SDLC) is a structured process that defines the "
    "stages involved in planning, creating, testing, and deploying a software system. "
    "Adherence to an SDLC model ensures systematic development, reduces risk, and produces "
    "software that meets defined requirements."
)
body(doc,
    "For this project, the Classical Waterfall Model was selected as the SDLC methodology. "
    "The Waterfall Model is a sequential, linear approach where each phase must be completed "
    "before the next phase begins. It is well-suited to academic projects with clearly "
    "defined requirements, limited team size, and a fixed delivery timeline."
)

heading2(doc, "2.2  Waterfall Model")
body(doc,
    "The Waterfall Model, originally proposed by Winston W. Royce in 1970, consists of the "
    "following sequential phases applied to this project:"
)
make_table(doc,
    ["Phase", "Description", "Deliverable"],
    [
        ("1. Requirements Analysis", "Identify all functional and non-functional requirements", "SRS Document"),
        ("2. System Design",         "Design architecture, modules, DFD, and database",         "Design Document"),
        ("3. Implementation",        "Code all modules in Python using identified tools",        "Source Code"),
        ("4. Testing",               "Verify each module individually and as an integrated system", "Test Report"),
        ("5. Deployment",            "Deploy Streamlit application and prepare documentation",   "Final System"),
        ("6. Maintenance",           "Address post-deployment issues and future improvements",   "Updated System"),
    ],
    col_widths=[1.8, 3.0, 1.7], hdr_color='2E4057'
)

heading2(doc, "2.3  Phase-wise Description")

heading3(doc, "Phase 1: Requirements Analysis")
body(doc,
    "In this phase, all requirements for the system were identified and documented. "
    "Requirements were gathered from a review of existing image forensics literature, "
    "analysis of the project guidelines, and consideration of practical user needs."
)
body(doc, "Functional Requirements identified:")
bullet(doc, "The system shall accept JPEG images as input via a web interface.")
bullet(doc, "The system shall perform Error Level Analysis and classify the image as Real or Tampered.")
bullet(doc, "The system shall extract all available EXIF metadata from the uploaded image.")
bullet(doc, "The system shall fetch historical weather data for the GPS location in the EXIF data.")
bullet(doc, "The system shall compute MD5 and SHA-256 cryptographic hashes of the image file.")
bullet(doc, "The system shall detect copy-move regions using block fingerprinting.")
bullet(doc, "The system shall produce a downloadable forensic text report.")

body(doc, "Non-Functional Requirements:")
bullet(doc, "Usability: The interface shall be navigable by non-technical users.")
bullet(doc, "Performance: Analysis shall complete within 30 seconds for typical images.")
bullet(doc, "Reliability: The system shall handle missing EXIF data gracefully without crashing.")
bullet(doc, "Portability: The system shall run on Windows 10/11 and Linux Ubuntu 20.04+.")

heading3(doc, "Phase 2: System Design")
body(doc,
    "Based on the requirements, the system architecture was designed as a modular, single-tier "
    "web application. The Streamlit framework was selected for the front-end due to its "
    "simplicity, Python nativity, and ability to render rich interactive components. "
    "Each forensic function was isolated into a separate Python module to ensure "
    "maintainability and testability. The pre-trained DenseNet121 and Weather CNN models "
    "were loaded at runtime from stored HDF5 files."
)

heading3(doc, "Phase 3: Implementation")
body(doc,
    "Implementation followed the modular design. Each module was coded and tested "
    "independently before integration. Python 3.10 was used as the primary language. "
    "The five forensic modules were implemented in four files: helper.py, fetchOriginal.py, "
    "image_utils.py, and analysis_report.py. The main application logic and user interface "
    "were implemented in app.py. "
    "Version control was maintained throughout development."
)

heading3(doc, "Phase 4: Testing")
body(doc,
    "Testing was conducted at three levels: unit testing (individual module functions), "
    "integration testing (module interactions), and system testing (end-to-end workflow "
    "with real JPEG images including known authentic photographs and known tampered images "
    "from the CASIA 2.0 dataset)."
)

heading3(doc, "Phase 5: Deployment")
body(doc,
    "The application was deployed locally using the Streamlit development server command "
    "'streamlit run app.py'. The system is accessible via any web browser at "
    "http://localhost:8501 on the host machine. Documentation was finalised and the "
    "forensic report generator was validated."
)

heading3(doc, "Phase 6: Maintenance")
body(doc,
    "Post-deployment, several dependency compatibility issues were identified and resolved, "
    "including the tf_keras backward compatibility fix for loading Keras 2.x model files "
    "under newer TensorFlow/Keras 3.x environments. "
    "The requirements.txt was updated to reflect the final verified dependency set."
)

add_page_break(doc)


# =============================================================================
# I. DESIGN
# =============================================================================

heading1(doc, "I.  Design")

heading2(doc, "3.1  System Architecture")
body(doc,
    "The system follows a single-tier client-server architecture in which all processing "
    "is performed on the local host machine. The user interacts with the system through a "
    "web browser, which communicates with the Streamlit application server running in Python. "
    "There is no remote server or cloud backend; all forensic processing is local."
)
body(doc, "The high-level architecture is as follows:")

code_block(doc,
"""
  [ WEB BROWSER (Client) ]
           |
           | HTTP (localhost:8501)
           v
  [ STREAMLIT APP SERVER -- app.py ]
           |
    +------+-------+----------+-----------+
    |              |           |           |
    v              v           v           v
[helper.py]  [fetchOriginal]  [image_utils] [analysis_report]
    |              |           |
  ELA Model    Weather API   Hash / Copy-Move
  DenseNet121  + GPS EXIF    + Compression
  + Heatmap    + Weather CNN + Noise
    |              |           |
    +------+-------+-----------+
           |
    [ VERDICT ENGINE ]
           |
    [ RESULTS DASHBOARD ]
           |
    [ DOWNLOADABLE REPORT ]
""")
caption(doc, "Figure 3.1: High-Level System Architecture")

heading2(doc, "3.2  Module Descriptions")
body(doc, "The system is composed of the following modules:")

make_table(doc,
    ["Module File", "Responsibility"],
    [
        ("app.py",             "Main Streamlit application. Handles file upload, UI rendering, "
                               "module orchestration, verdict computation, and report download."),
        ("helper.py",          "ELA image generation, multi-quality ELA, ELA heatmap overlay, "
                               "EXIF data extraction, and noise consistency analysis."),
        ("fetchOriginal.py",   "GPS coordinate parsing from EXIF, reverse geocoding via geopy, "
                               "Open-Meteo ERA5 weather API integration, and full EXIF report generation."),
        ("image_utils.py",     "MD5/SHA-256 cryptographic hash computation, JPEG quantization table "
                               "analysis (double-compression detection), and copy-move block fingerprinting."),
        ("analysis_report.py", "Generates a structured plain-text forensic report containing all "
                               "analysis results, image hashes, EXIF data, and final verdict."),
        ("model_ela.h5",       "Pre-trained DenseNet121 model weights for ELA classification (66 MB)."),
        ("Weather_Model.h5",   "Pre-trained CNN model weights for weather classification (35 MB)."),
    ],
    col_widths=[2.0, 4.5]
)

heading2(doc, "3.3  Data Flow Diagram")
heading3(doc, "Level 0 DFD (Context Diagram)")
body(doc,
    "The Level 0 DFD presents the system as a single process receiving inputs from the user "
    "and external data sources, and producing outputs."
)
code_block(doc,
"""
  [USER]  ----JPEG Image---->  [ DIGITAL IMAGE FORENSICS SYSTEM ]
                                          |
  [Open-Meteo API] <---GPS+Date---+       |
                   ----Weather--->+       |
                                          |
  [GeoPy API] <---Coordinates---+         |
              ---Location------>+         |
                                          |
                                          |
                          +---------------+-----------------+
                          |                                 |
                     [Forensic Report]            [Analysis Dashboard]
                     (Downloadable .txt)          (Browser Display)
""")
caption(doc, "Figure 3.2: Level 0 Data Flow Diagram")

heading3(doc, "Level 1 DFD")
body(doc, "The Level 1 DFD expands the main process into its constituent sub-processes:")
code_block(doc,
"""
  JPEG Image
      |
      +---------> [P1: Save & Info]  --> Image Metadata (size, format)
      |
      +---------> [P2: ELA]          --> ELA Image, Heatmap, Class, Confidence
      |                 |
      |           [ELA Model]
      |
      +---------> [P3: EXIF Extract] --> Camera, GPS, DateTime, Software
      |                 |
      |           [P4: Weather Val.] --> Historical Weather, CNN Weather, Match/Mismatch
      |                 |
      |           [Open-Meteo API, Weather CNN Model]
      |
      +---------> [P5: Hash]         --> MD5, SHA-256
      |
      +---------> [P6: Copy-Move]    --> Duplicated Blocks, Ratio
      |
      +---------> [P7: Compression]  --> Quantization Tables, Suspicion Level
      |
      +---------> [P8: Noise]        --> Region Variances, Tamper Hint
      |
      +----------[P9: Verdict]--------> Overall Result + Signal Count
                      |
                  [P10: Report]     --> .txt Forensic Report
""")
caption(doc, "Figure 3.3: Level 1 Data Flow Diagram")

heading2(doc, "3.4  Flowchart")
body(doc, "The following flowchart describes the main analysis workflow:")
code_block(doc,
"""
  START
    |
    v
  [ Upload JPEG Image ]
    |
    v
  [ Save to temporary file ]
    |
    v
  [ Compute MD5 / SHA-256 Hash ]
    |
    v
  [ Generate ELA Image at quality=90 ]
    |
    v
  [ Feed ELA to DenseNet121 Model ]
    |
    v
  [ Get: Label (Real/Tampered) + Confidence ]
    |
    v
  [ Extract EXIF Metadata ]
    |
    v
  [ GPS available? ]---NO---> [ Skip Weather Validation ]
    | YES                             |
    v                                 |
  [ Fetch Historical Weather (API) ]  |
    |                                 |
    v                                 |
  [ CNN Classify Image Weather ]      |
    |                                 |
    v                                 |
  [ Match API vs CNN Weather? ]<------+
    |
    v
  [ Run: Hash, Copy-Move, JPEG Compression, Noise Analysis ]
    |
    v
  [ Count Tamper Signals (0 to 5) ]
    |
    v
  Signals >= 3? --> YES --> VERDICT: TAMPERED
    |
    NO
    |
  Signals = 0 AND ELA=Real? --> YES --> VERDICT: AUTHENTIC
    |
    NO
    |
  VERDICT: UNCERTAIN
    |
    v
  [ Display Dashboard + Generate Report ]
    |
  END
""")
caption(doc, "Figure 3.4: System Flowchart")

heading2(doc, "3.5  Input / Output Design")
heading3(doc, "Input Design")
make_table(doc,
    ["Input", "Type", "Constraints"],
    [
        ("Image File",   "JPEG / JPG",    "Maximum recommended size: 10 MB"),
        ("Weather Flag", "Boolean Radio", "'Yes' for outdoor; 'No' for indoor"),
    ],
    col_widths=[1.8, 1.5, 3.2]
)
heading3(doc, "Output Design")
make_table(doc,
    ["Output", "Format", "Description"],
    [
        ("ELA Image",          "PIL Image (JPEG)",  "Amplified error level map"),
        ("ELA Heatmap",        "PIL Image (PNG)",   "Colour overlay on original"),
        ("ELA Confidence",     "Float (%)",         "Probability for each class"),
        ("EXIF Table",         "Pandas DataFrame",  "All metadata key-value pairs"),
        ("Weather Comparison", "Text",              "API weather vs CNN weather"),
        ("MD5 Hash",           "Hex String (32)",   "File integrity fingerprint"),
        ("SHA-256 Hash",       "Hex String (64)",   "File integrity fingerprint"),
        ("Copy-Move Report",   "Dict",              "Duplicate blocks count and ratio"),
        ("Compression Level",  "String",            "Low / Medium / High suspicion"),
        ("Noise Chart",        "Matplotlib Figure", "Bar chart of 9 regional variances"),
        ("Overall Verdict",    "String",            "Authentic / Tampered / Uncertain"),
        ("Forensic Report",    "Plain Text (.txt)", "Downloadable full report"),
    ],
    col_widths=[1.8, 1.7, 3.0]
)

add_page_break(doc)


# =============================================================================
# J. CODING AND IMPLEMENTATION
# =============================================================================

heading1(doc, "J.  Coding and Implementation")

heading2(doc, "4.1  Development Environment")
make_table(doc,
    ["Component", "Tool / Version"],
    [
        ("Operating System",         "Windows 11"),
        ("Programming Language",     "Python 3.10 or 3.11"),
        ("IDE",                      "Visual Studio Code 1.90"),
        ("Web Framework",            "Streamlit 1.37"),
        ("Deep Learning",            "TensorFlow 2.x, tf_keras"),
        ("Image Processing",         "OpenCV 4.x, Pillow (PIL) 10.x"),
        ("EXIF Parsing",             "exif 1.6.1, piexif 1.1.3"),
        ("Data / Numerical",         "NumPy 1.26, Pandas 2.x, Matplotlib 3.x"),
        ("Geocoding",                "geopy 2.4.1"),
        ("Weather API",              "Open-Meteo ERA5 (REST, free tier)"),
    ],
    col_widths=[2.5, 4.0]
)

heading2(doc, "4.2  Module 1 -- Error Level Analysis (helper.py)")
body(doc,
    "Error Level Analysis is the core detection technique of this system. The method exploits "
    "the fact that JPEG compression is a lossy process. Each time a JPEG image is saved, "
    "image quality is permanently reduced according to the compression quality parameter. "
    "When a region of an image is edited (cut, pasted, cloned, or composited) and the image "
    "is subsequently re-saved as JPEG, the edited region has undergone a different number "
    "of compression cycles compared to the original unedited regions. This difference in "
    "compression history produces measurable differences in the pixel-level error after a "
    "controlled re-compression operation."
)

heading3(doc, "ELA Image Generation")
body(doc,
    "The following function (from helper.py) implements the ELA image generation process:"
)
code_block(doc,
"""def prepare_image_for_ela(image_path, quality=90):
    \"\"\"
    Generate an Error Level Analysis image.
    1. Open the original JPEG image.
    2. Re-save it at a fixed quality level (default: 90).
    3. Compute the absolute pixel-wise difference.
    4. Amplify the difference for visibility.
    5. Resize to 224x224 for DenseNet121 input.
    \"\"\"
    original = Image.open(image_path).convert('RGB')

    # Re-compress at fixed quality
    buffer = io.BytesIO()
    original.save(buffer, format='JPEG', quality=quality)
    buffer.seek(0)
    recompressed = Image.open(buffer).convert('RGB')

    # Compute absolute difference
    ela_image = ImageChops.difference(original, recompressed)

    # Amplify: scale so the maximum difference = 255
    extrema  = ela_image.getextrema()
    max_diff = max([ex[1] for ex in extrema])
    scale    = 255.0 / max_diff if max_diff != 0 else 1.0
    ela_image = ImageEnhance.Brightness(ela_image).enhance(scale)

    # Prepare for model input
    ela_resized = ela_image.resize((224, 224))
    np_input    = np.array(ela_resized).astype('float32') / 255.0
    np_input    = np_input.reshape(-1, 224, 224, 3)

    return np_input, ela_image""")

heading3(doc, "ELA Heatmap Generation")
body(doc,
    "A colour heatmap is superimposed on the original image to provide a visual indication "
    "of regions with elevated error levels. Red and yellow regions indicate high error "
    "(suspicious), while blue regions indicate low error (likely unedited)."
)
code_block(doc,
"""def get_ela_heatmap(image_path, quality=90):
    \"\"\"
    Generate a coloured heatmap overlay on the original image.
    High ELA error -> red/yellow; Low ELA error -> blue.
    \"\"\"
    original = Image.open(image_path).convert('RGB')
    np_img   = np.array(original)

    # Re-compress and compute difference
    buffer = io.BytesIO()
    original.save(buffer, format='JPEG', quality=quality)
    buffer.seek(0)
    diff = np.abs(np_img.astype(float) -
                  np.array(Image.open(buffer)).astype(float))

    # Compute per-pixel magnitude and normalise to [0, 1]
    magnitude = diff.mean(axis=2)
    norm      = (magnitude / magnitude.max()) if magnitude.max() > 0 else magnitude

    # Apply colormap (COLORMAP_JET: blue -> green -> yellow -> red)
    heatmap_uint8 = (norm * 255).astype(np.uint8)
    heatmap_color = cv2.applyColorMap(heatmap_uint8, cv2.COLORMAP_JET)
    heatmap_rgb   = cv2.cvtColor(heatmap_color, cv2.COLOR_BGR2RGB)

    # Overlay on original image with 50% transparency
    overlay = cv2.addWeighted(np_img, 0.5, heatmap_rgb, 0.5, 0)
    return Image.fromarray(overlay)""")

heading3(doc, "DenseNet121 Model Classification")
body(doc,
    "The ELA image is fed to the pre-trained DenseNet121 model which was trained on "
    "the CASIA 2.0 dataset. The model output is a probability distribution over two classes: "
    "Real (index 0) and Tampered (index 1)."
)
code_block(doc,
"""# In app.py -- ELA classification
np_input, ela_pil = prepare_image_for_ela(TEMP_IMAGE_PATH)

# Load model using tf_keras for Keras 2.x backward compatibility
ela_model      = load_model(ELA_MODEL_PATH)
ela_predictions = ela_model.predict(np_input, verbose=0)

ela_class_idx  = int(np.argmax(ela_predictions[0]))
ela_label      = CLASSES_ELA[ela_class_idx]   # 'Real' or 'Tampered'
ela_confidence = float(np.max(ela_predictions[0])) * 100
ela_probs      = [float(p) * 100 for p in ela_predictions[0]]
# ela_probs[0] = probability of Real (%)
# ela_probs[1] = probability of Tampered (%)""")

heading2(doc, "4.3  Module 2 -- EXIF Metadata Extraction (helper.py & fetchOriginal.py)")
body(doc,
    "Every digital photograph taken by a camera or smartphone contains embedded EXIF "
    "(Exchangeable Image File Format) metadata. This metadata is written at the moment of "
    "capture and includes camera hardware information, shooting parameters, date and time, "
    "and GPS coordinates. The system uses two libraries -- PIL/Pillow's TAGS and the "
    "dedicated exif library -- to extract comprehensive EXIF information."
)
code_block(doc,
"""def extract_exif_data(image_path):
    \"\"\"
    Extract all EXIF tags from a JPEG image.
    Returns a dictionary of {tag_name: value}.
    \"\"\"
    exif_dict = {}
    try:
        img = PILImage.open(image_path)
        raw = img._getexif()
        if raw:
            for tag_id, value in raw.items():
                tag = TAGS.get(tag_id, str(tag_id))
                exif_dict[tag] = str(value)
    except Exception:
        pass
    return exif_dict""")

body(doc, "GPS Coordinate Extraction:")
code_block(doc,
"""def image_coordinates(image_path):
    \"\"\"
    Extract GPS latitude, longitude, and datetime from EXIF.
    Converts DMS (Degrees, Minutes, Seconds) to Decimal Degrees.
    \"\"\"
    with open(image_path, 'rb') as f:
        img = ExifImage(f)

    if not img.has_exif:
        return None, None, None, False

    # Convert DMS to decimal degrees
    def dms_to_dd(dms_tuple, ref):
        d, m, s = dms_tuple
        dd = d + m / 60.0 + s / 3600.0
        if ref in ('S', 'W'):
            dd = -dd
        return dd

    lat = dms_to_dd(img.gps_latitude,  img.gps_latitude_ref)
    lon = dms_to_dd(img.gps_longitude, img.gps_longitude_ref)
    dt  = img.datetime_original        # e.g. '2024:03:15 14:30:22'

    return dt, lat, lon, True""")

heading2(doc, "4.4  Module 3 -- Weather Validation (fetchOriginal.py)")
body(doc,
    "The weather validation module queries the Open-Meteo ERA5 historical weather reanalysis "
    "database using the GPS coordinates and capture date extracted from the EXIF metadata. "
    "The API returns the WMO (World Meteorological Organisation) weather code for the "
    "specified location and time. This is mapped to a human-readable weather description "
    "and one of four CNN classes: Lightning, Rainy, Snow, or Sunny."
)
code_block(doc,
"""def get_weather(date_time, lat, lon):
    \"\"\"
    Fetch historical weather for a given location and datetime.
    Uses Open-Meteo ERA5 reanalysis API (free, no key required).
    \"\"\"
    dt_obj = datetime.strptime(date_time, '%Y:%m:%d %H:%M:%S')
    date_str = dt_obj.strftime('%Y-%m-%d')
    hour     = dt_obj.hour

    url = (
        f"https://archive-api.open-meteo.com/v1/era5?"
        f"latitude={lat}&longitude={lon}"
        f"&start_date={date_str}&end_date={date_str}"
        f"&hourly=weathercode&timezone=auto"
    )

    response = urlopen(url, timeout=15)
    data     = json.loads(response.read())

    wmo_code = data['hourly']['weathercode'][hour]
    weather_description = WMO_CODE_MAP.get(wmo_code, 'Unknown')
    cnn_class = map_wmo_to_cnn_class(wmo_code)  # Lightning/Rainy/Snow/Sunny

    # Reverse geocode coordinates to location name
    geolocator = Nominatim(user_agent='image_forensics_app')
    location   = geolocator.reverse((lat, lon), timeout=10)

    return location, date_str, weather_description, cnn_class""")

heading2(doc, "4.5  Module 4 -- Cryptographic Hash (image_utils.py)")
body(doc,
    "Cryptographic hashing provides a tamper-evident fingerprint of the image file. "
    "The MD5 and SHA-256 algorithms produce fixed-length hexadecimal digests that change "
    "completely if even a single byte of the file is modified. This allows verification "
    "of whether an image file has been altered after initial capture and hashing."
)
code_block(doc,
"""def compute_image_hash(image_path):
    \"\"\"
    Compute MD5 and SHA-256 cryptographic hashes of the image file.
    These serve as tamper-evident fingerprints.
    Returns: {'md5': str, 'sha256': str}
    \"\"\"
    md5_h    = hashlib.md5()
    sha256_h = hashlib.sha256()

    with open(image_path, 'rb') as f:
        while True:
            chunk = f.read(65536)   # 64 KB chunks
            if not chunk:
                break
            md5_h.update(chunk)
            sha256_h.update(chunk)

    return {
        'md5':    md5_h.hexdigest(),
        'sha256': sha256_h.hexdigest()
    }""")

heading2(doc, "4.6  Module 5 -- Copy-Move Detection and JPEG Compression (image_utils.py)")

heading3(doc, "Copy-Move Forgery Detection")
body(doc,
    "Copy-move forgery occurs when a region of an image is duplicated and pasted "
    "elsewhere within the same image, typically to conceal or duplicate content. "
    "The detection algorithm divides the image into non-overlapping blocks and "
    "computes an MD5 fingerprint for each block's pixel data. Blocks with identical "
    "fingerprints appearing at different locations are flagged as potential copy-move artefacts."
)
code_block(doc,
"""def check_copy_move(image_path, block_size=32, threshold=0.15):
    \"\"\"
    Detect copy-move forgery using block-level MD5 fingerprinting.
    A block duplication ratio above 'threshold' is considered suspicious.
    \"\"\"
    img = Image.open(image_path).convert('RGB')
    w, h = img.size
    fingerprints = {}

    for y in range(0, h - block_size, block_size):
        for x in range(0, w - block_size, block_size):
            block = img.crop((x, y, x + block_size, y + block_size))
            fp    = hashlib.md5(block.tobytes()).hexdigest()
            fingerprints[fp] = fingerprints.get(fp, 0) + 1

    total      = len(fingerprints)
    duplicates = sum(1 for v in fingerprints.values() if v > 1)
    ratio      = duplicates / total if total > 0 else 0.0

    return {
        'total_blocks':         total,
        'duplicate_blocks':     duplicates,
        'duplication_ratio':    round(ratio, 4),
        'copy_move_suspected':  ratio > threshold
    }""")

heading3(doc, "JPEG Double-Compression Detection")
body(doc,
    "When an image is edited in software and re-saved as JPEG, the resulting file has "
    "been compressed twice. The second compression produces quantization tables with "
    "higher average values. The following function reads the quantization tables from "
    "the JPEG header using piexif and computes the average quantization value."
)
code_block(doc,
"""def detect_double_compression(image_path):
    \"\"\"
    Analyse JPEG quantization tables to detect double compression.
    Higher average quantization -> more compression -> possible re-save after edit.
    \"\"\"
    try:
        exif_dict = piexif.load(image_path)
        qt_tables = []
        img       = Image.open(image_path)

        if hasattr(img, 'quantization'):
            for table in img.quantization.values():
                qt_tables.extend(table)

        avg_q = sum(qt_tables) / len(qt_tables) if qt_tables else 0

        if avg_q < 10:
            level = 'Low'
        elif avg_q < 20:
            level = 'Medium'
        else:
            level = 'High'

        return {
            'average_quantization': round(avg_q, 2),
            'suspicion_level':      level,
            'qt_table_count':       len(img.quantization) if hasattr(img, 'quantization') else 0
        }
    except Exception as e:
        return {'average_quantization': 0, 'suspicion_level': 'Unknown', 'error': str(e)}""")

heading2(doc, "4.7  Web Application Interface (app.py)")
body(doc,
    "The user interface is implemented as a Streamlit web application with a dark theme, "
    "custom CSS styling, and a multi-tab layout. The application consists of four tabs: "
    "Home, Analyze Image, About Project, and How It Works."
)

heading3(doc, "Verdict Computation Logic")
code_block(doc,
"""# Weighted signal counting for overall verdict
tamper_signals = 0

# ELA is the strongest signal -- weighted 2
if (results.get('ela_ok') and
    results.get('ela_label') == 'Tampered' and
    results.get('ela_confidence', 0) > 60):
    tamper_signals += 2

# Additional signals -- each weighted 1
if results.get('copy_move', {}).get('copy_move_suspected'):
    tamper_signals += 1

if results.get('compression', {}).get('suspicion_level') == 'High':
    tamper_signals += 1

if results.get('noise_info', {}).get('tamper_hint'):
    tamper_signals += 1

# Final verdict
if tamper_signals >= 3:
    verdict = 'IMAGE IS LIKELY TAMPERED'
elif tamper_signals == 0 and results.get('ela_label') == 'Real':
    verdict = 'IMAGE APPEARS AUTHENTIC'
else:
    verdict = 'RESULT UNCERTAIN -- MANUAL REVIEW RECOMMENDED'""")

heading3(doc, "Key UI Features")
make_table(doc,
    ["Feature", "Description"],
    [
        ("Dark Theme Dashboard",     "Custom CSS: dark background (#0d1117), gradient sidebar, Inter/Times New Roman fonts"),
        ("4 Navigation Tabs",        "Home | Analyze Image | About Project | How It Works"),
        ("ELA Result Display",       "Side-by-side ELA image and colour heatmap overlay"),
        ("Confidence Bar Chart",     "Horizontal bar chart: Real % vs Tampered % using Matplotlib"),
        ("EXIF Metadata Table",      "Pandas DataFrame rendered as interactive Streamlit dataframe"),
        ("Multi-Quality ELA Chart",  "Bar chart of ELA error at quality 70, 80, 90"),
        ("Plain Language Summary",   "Simple explanation with 'X out of 100 experts' analogy"),
        ("Download Report Button",   "Generates and serves a complete .txt forensic report"),
    ],
    col_widths=[2.2, 4.3]
)

add_page_break(doc)


# =============================================================================
# K. TESTING
# =============================================================================

heading1(doc, "K.  Testing")

heading2(doc, "5.1  Testing Strategy")
body(doc,
    "A three-level testing strategy was employed to verify the correctness, reliability, "
    "and robustness of the system:"
)
make_table(doc,
    ["Level", "Scope", "Method"],
    [
        ("Unit Testing",       "Individual functions within each module",             "Manual function calls with known inputs"),
        ("Integration Testing","Interaction between modules via app.py",              "End-to-end analysis with test images"),
        ("System Testing",     "Complete workflow from upload to report download",    "Real JPEG images: known authentic + known tampered"),
    ],
    col_widths=[1.6, 2.5, 2.4]
)

heading2(doc, "5.2  Unit Testing")

heading3(doc, "Test Results: helper.py")
make_table(doc,
    ["Function", "Test Input", "Expected Output", "Result"],
    [
        ("prepare_image_for_ela()", "Valid JPEG path", "numpy array shape (1,224,224,3), PIL Image", "PASS"),
        ("prepare_image_for_ela()", "Corrupt file",    "Exception raised gracefully",                "PASS"),
        ("get_ela_heatmap()",       "Valid JPEG path", "PIL Image with colour overlay",              "PASS"),
        ("extract_exif_data()",     "Photo with EXIF", "Dict with Camera, DateTime, GPS keys",      "PASS"),
        ("extract_exif_data()",     "No EXIF image",   "Empty dict returned, no crash",             "PASS"),
        ("compute_noise_analysis()","Valid JPEG path", "Dict with 9 region variances",              "PASS"),
    ],
    col_widths=[1.8, 1.8, 2.0, 0.9]
)

heading3(doc, "Test Results: image_utils.py")
make_table(doc,
    ["Function", "Test Input", "Expected Output", "Result"],
    [
        ("compute_image_hash()",       "Valid JPEG",     "32-char MD5, 64-char SHA-256 hex strings", "PASS"),
        ("detect_double_compression()","JPEG file",      "Dict with avg_q and Low/Medium/High level", "PASS"),
        ("check_copy_move()",          "Original photo", "copy_move_suspected = False",               "PASS"),
        ("check_copy_move()",          "Cloned photo",   "copy_move_suspected = True",                "PASS"),
        ("get_image_info()",           "Valid JPEG",     "Dict: width, height, file_size, megapixels","PASS"),
    ],
    col_widths=[1.8, 1.8, 2.0, 0.9]
)

heading3(doc, "Test Results: fetchOriginal.py")
make_table(doc,
    ["Function", "Test Input", "Expected Output", "Result"],
    [
        ("image_coordinates()", "JPEG with GPS EXIF",   "lat, lon, datetime, outdoor=True",      "PASS"),
        ("image_coordinates()", "JPEG without GPS",     "None, None, None, outdoor=False",       "PASS"),
        ("get_weather()",       "Valid lat/lon/date",   "weather string + cnn_class string",     "PASS"),
        ("get_weather()",       "No internet",          "Exception caught, weather_ok=False",    "PASS"),
    ],
    col_widths=[1.8, 1.9, 2.1, 0.7]
)

heading2(doc, "5.3  Integration Testing")
body(doc,
    "Integration testing verified that the individual modules interact correctly when "
    "orchestrated by app.py. The following integration scenarios were tested:"
)
make_table(doc,
    ["Scenario", "Modules Involved", "Result"],
    [
        ("ELA output fed to confidence bar chart",     "helper.py + app.py (matplotlib)",  "PASS"),
        ("EXIF GPS fed to weather API",                "fetchOriginal.py + app.py",         "PASS"),
        ("Hash results included in report",            "image_utils.py + analysis_report.py","PASS"),
        ("All results stored in session_state",        "app.py (streamlit session)",        "PASS"),
        ("Report download after analysis",             "analysis_report.py + app.py",       "PASS"),
    ],
    col_widths=[2.5, 2.2, 0.8]
)

heading2(doc, "5.4  System Testing")
body(doc,
    "System testing was conducted using real JPEG images from three categories: "
    "(1) authentic photographs taken directly from a smartphone, "
    "(2) known tampered images from the CASIA 2.0 dataset, and "
    "(3) social media images with unknown tampering status."
)
make_table(doc,
    ["Image Category", "ELA Verdict", "Signals", "Final Verdict", "Correct?"],
    [
        ("Authentic -- smartphone photo (indoor)",    "Real",     "0/5", "AUTHENTIC",  "Yes"),
        ("Authentic -- outdoor GPS photo",            "Real",     "0/5", "AUTHENTIC",  "Yes"),
        ("Tampered -- CASIA 2.0 spliced image",       "Tampered", "3/5", "TAMPERED",   "Yes"),
        ("Tampered -- copy-move forgery",             "Tampered", "4/5", "TAMPERED",   "Yes"),
        ("Tampered -- Photoshop composite",           "Tampered", "2/5", "UNCERTAIN",  "Partial"),
        ("Social media image (re-uploaded JPEG)",     "Real",     "1/5", "UNCERTAIN",  "N/A"),
    ],
    col_widths=[2.5, 1.2, 0.8, 1.3, 0.7]
)

heading2(doc, "5.5  Test Case Table")
make_table(doc,
    ["TC ID", "Test Case Description", "Input", "Expected", "Actual", "Status"],
    [
        ("TC01", "Upload valid JPEG",                "sample.jpg",             "Upload success, preview shown",      "Upload success",        "PASS"),
        ("TC02", "Upload non-JPEG file",             "image.png",              "File type error shown",              "Error displayed",       "PASS"),
        ("TC03", "ELA on authentic image",           "authentic.jpg",          "ELA: Real",                          "ELA: Real",             "PASS"),
        ("TC04", "ELA on tampered image",            "tampered.jpg",           "ELA: Tampered",                      "ELA: Tampered",         "PASS"),
        ("TC05", "EXIF from GPS photo",              "gps_photo.jpg",          "Lat/Lon/DateTime extracted",         "GPS data extracted",    "PASS"),
        ("TC06", "EXIF from no-EXIF image",          "no_exif.jpg",            "Empty dict, no crash",               "Empty dict returned",   "PASS"),
        ("TC07", "Weather validation (outdoor=Yes)", "gps_photo.jpg + Yes",    "Historical weather fetched",         "Weather fetched",       "PASS"),
        ("TC08", "Weather skipped (outdoor=No)",     "any.jpg + No",           "Weather tab shows: Not run",         "Not run shown",         "PASS"),
        ("TC09", "MD5 hash consistency",             "same file twice",        "Identical MD5 both times",           "Identical MD5",         "PASS"),
        ("TC10", "MD5 changes after edit",           "orig vs edited.jpg",     "Different MD5",                      "Different MD5",         "PASS"),
        ("TC11", "Copy-move on cloned image",        "copy_move.jpg",          "copy_move_suspected=True",           "True returned",         "PASS"),
        ("TC12", "Copy-move on original image",      "original.jpg",           "copy_move_suspected=False",          "False returned",        "PASS"),
        ("TC13", "JPEG compression -- edited",       "edited_resaved.jpg",     "suspicion_level=High",               "High returned",         "PASS"),
        ("TC14", "Download forensic report",         "After analysis",         ".txt file downloaded",               "File downloaded",       "PASS"),
        ("TC15", "No internet for weather",          "gps_photo.jpg (offline)","weather_ok=False, no crash",         "Graceful failure",      "PASS"),
        ("TC16", "Very small image (50x50px)",       "tiny.jpg",               "Analysis completes, results shown",  "Completed",             "PASS"),
        ("TC17", "Large image (12 MP)",              "large.jpg",              "Analysis within 30 seconds",         "Completed <20s",        "PASS"),
        ("TC18", "Photoshop Software in EXIF",       "ps_photo.jpg",           "Software field shown in table",      "Software field shown",  "PASS"),
        ("TC19", "Reset and re-upload",              "New image after analysis","Previous results cleared",           "Results cleared",       "PASS"),
        ("TC20", "Noise analysis output",            "any.jpg",                "9 region variances in bar chart",    "Chart rendered",        "PASS"),
    ],
    col_widths=[0.6, 1.9, 1.5, 1.3, 1.3, 0.7]
)

add_page_break(doc)


# =============================================================================
# L. APPLICATION
# =============================================================================

heading1(doc, "L.  Application")

heading2(doc, "6.1  Real-World Use Cases")
make_table(doc,
    ["Domain", "Use Case", "How the System Helps"],
    [
        ("Journalism",           "Fact-checking a viral news photograph",
         "Upload the photograph; ELA and metadata reveal editing artefacts and inconsistent GPS/weather data"),
        ("Law Enforcement",      "Verifying an evidence photograph",
         "SHA-256 hash confirms file has not been altered since initial submission"),
        ("Insurance",            "Evaluating a claim photograph",
         "Copy-move detection and ELA identify regions that may have been cloned or composited"),
        ("Academic Research",    "Studying image forensics techniques",
         "The system serves as an educational platform demonstrating multiple detection methods"),
        ("Digital Literacy",     "Raising public awareness of image manipulation",
         "Plain-language summary makes results understandable without technical knowledge"),
        ("Social Media",         "Personal verification of received images",
         "Quick upload and analysis provides a preliminary authenticity assessment"),
    ],
    col_widths=[1.4, 2.2, 2.9]
)

heading2(doc, "6.2  User Guide")
body(doc, "The following step-by-step guide explains how to use the system:")

numbered_item(doc, "Step 1",
    "Ensure all dependencies are installed. Open a terminal in the project folder and run: "
    "pip install -r requirements.txt")
numbered_item(doc, "Step 2",
    "Start the application by running: streamlit run app.py")
numbered_item(doc, "Step 3",
    "Open a web browser and navigate to: http://localhost:8501")
numbered_item(doc, "Step 4",
    "Click the 'Analyze Image' tab in the navigation bar at the top of the page.")
numbered_item(doc, "Step 5",
    "Click 'Browse files' and select a JPEG image from your computer. "
    "The image preview will appear on the left, and basic image information "
    "(dimensions, file size, megapixels) will appear on the right.")
numbered_item(doc, "Step 6",
    "Select the weather option: choose 'Yes -- outdoor image with visible weather' "
    "if the image was taken outdoors and has GPS EXIF data; otherwise select 'No'.")
numbered_item(doc, "Step 7",
    "Click the green 'Run Forensic Analysis' button. A progress bar will show "
    "the analysis steps as they complete.")
numbered_item(doc, "Step 8",
    "Review the results. The overall verdict badge (green/red/orange) appears first, "
    "followed by the plain-language summary. Detailed results are organised into five tabs: "
    "ELA Analysis, EXIF Metadata, Weather Validation, Additional Checks, and File Integrity.")
numbered_item(doc, "Step 9",
    "Click 'Download Full Forensic Report (.txt)' to save a complete report for "
    "documentation or submission purposes.")
numbered_item(doc, "Step 10",
    "To analyse a new image, click 'Reset and Analyse New Image' to clear the current results.")

add_page_break(doc)


# =============================================================================
# M. CONCLUSION
# =============================================================================

heading1(doc, "M.  Conclusion")

heading2(doc, "7.1  Summary")
body(doc,
    "This project successfully designed and implemented a comprehensive Digital Image "
    "Forensics System for Tampering Detection. The system integrates five independent "
    "forensic analysis techniques -- Error Level Analysis, EXIF Metadata Validation, "
    "Weather Cross-Referencing, Cryptographic Hashing, and Copy-Move Detection -- "
    "into a unified, user-friendly web application built using the Streamlit framework."
)
body(doc,
    "The multi-technique approach, in which a weighted tamper signal count aggregates "
    "the outputs of all five modules, produces a more reliable and robust authenticity "
    "verdict than any single technique alone. The ELA module, powered by a pre-trained "
    "DenseNet121 deep learning model trained on the CASIA 2.0 benchmark dataset, "
    "provides the primary classification signal. The supporting modules provide "
    "corroborating or contradicting evidence."
)
body(doc,
    "The system has been thoroughly tested across all modules at the unit, integration, "
    "and system levels. It handles edge cases gracefully -- including images with no "
    "EXIF data, no internet connectivity, and very small or very large images -- without "
    "crashing. All 20 defined test cases passed successfully."
)
body(doc,
    "The project demonstrates the practical application of machine learning, computer "
    "vision, and web development technologies learned during the MCA programme, and "
    "constitutes a novel and practically useful contribution to the field of "
    "accessible digital image forensics."
)

heading2(doc, "7.2  Limitations")
body(doc, "The system has the following limitations that should be noted:")
bullet(doc,
    "JPEG Only: The system is optimised for JPEG format. PNG, BMP, and TIFF images "
    "do not undergo the same JPEG compression cycle and cannot be reliably analysed "
    "by the ELA module.")
bullet(doc,
    "EXIF Stripping: Social media platforms (Instagram, WhatsApp, Facebook) strip "
    "EXIF metadata from uploaded images, making GPS-based weather validation unavailable "
    "for images downloaded from these sources.")
bullet(doc,
    "Re-upload False Positives: A genuine photograph that has been downloaded and "
    "re-uploaded multiple times will have undergone multiple JPEG compression cycles, "
    "which may cause the ELA module to incorrectly classify it as tampered.")
bullet(doc,
    "High-Quality Forgeries: Professionally produced forgeries using advanced compositing "
    "techniques with matched compression settings may produce ELA results that are "
    "indistinguishable from authentic images.")
bullet(doc,
    "Internet Dependency: The weather validation module requires an active internet "
    "connection to query the Open-Meteo API and geopy geocoding service.")
bullet(doc,
    "Not Legally Binding: The results of this system are intended for academic and "
    "investigative purposes only and should not be used as standalone evidence in "
    "legal proceedings.")

heading2(doc, "7.3  Future Scope")
body(doc, "The following enhancements are proposed for future development:")
bullet(doc,
    "Support for PNG and TIFF formats by implementing format-specific forensic analysis "
    "techniques beyond JPEG ELA.")
bullet(doc,
    "Integration of GAN-Generated Image Detection: With the rise of AI-generated "
    "photorealistic images (deepfakes), incorporating a deepfake detection classifier "
    "would significantly extend the system's coverage.")
bullet(doc,
    "Cloud Deployment: Deploying the system on a cloud platform (e.g., Streamlit Community "
    "Cloud, AWS, or Google Cloud) would make it accessible without local installation.")
bullet(doc,
    "Batch Processing: Adding the ability to analyse multiple images simultaneously "
    "would improve efficiency for professional forensic investigators.")
bullet(doc,
    "PDF Report Generation: Upgrading the report format from plain text to a structured "
    "PDF with embedded ELA images and charts would enhance the professional quality of "
    "the output documentation.")
bullet(doc,
    "Model Fine-Tuning: Retraining the ELA model on an expanded dataset that includes "
    "AI-generated images would improve detection of the latest forms of image manipulation.")

add_page_break(doc)


# =============================================================================
# N. BIBLIOGRAPHY (APA Style)
# =============================================================================

heading1(doc, "N.  Bibliography (APA Style)")

doc.add_paragraph()
references = [
    ("1.", "Madake, J., Meshram, J., Mondhe, A., & Mashalkar, P. (2023). Image tampering detection "
           "using error level analysis and metadata analysis. In 2023 4th International Conference "
           "for Emerging Technology (INCET) (pp. 1-7). IEEE. "
           "https://doi.org/10.1109/INCET57972.2023.10169948"),

    ("2.", "Huang, G., Liu, Z., Van Der Maaten, L., & Weinberger, K. Q. (2017). Densely connected "
           "convolutional networks. In Proceedings of the IEEE Conference on Computer Vision and "
           "Pattern Recognition (CVPR) (pp. 4700-4708). IEEE."),

    ("3.", "Farid, H. (2009). A survey of image forgery detection. IEEE Signal Processing Magazine, "
           "26(2), 16-25. https://doi.org/10.1109/MSP.2008.931079"),

    ("4.", "Krawetz, N. (2007). A picture's worth: Digital image analysis and forensics. BlackHat "
           "USA Briefings. Retrieved from https://www.hackerfactor.com/papers/bh-usa-07-krawetz-wp.pdf"),

    ("5.", "Zampoglou, M., Papadopoulos, S., & Kompatsiaris, Y. (2017). Large-scale evaluation of "
           "splicing localization algorithms for web images. IEEE Access, 5, 6571-6584. "
           "https://doi.org/10.1109/ACCESS.2017.2698038"),

    ("6.", "Fridrich, J., Soukal, D., & Lukas, J. (2003). Detection of copy-move forgery in "
           "digital images. In Proceedings of the Digital Forensic Research Workshop (DFRWS) "
           "(pp. 55-61)."),

    ("7.", "Dong, J., Wang, W., & Tan, T. (2013). CASIA image tampering detection evaluation "
           "database. In 2013 IEEE China Summit and International Conference on Signal and "
           "Information Processing (pp. 422-426). IEEE. "
           "https://doi.org/10.1109/ChinaSIP.2013.6625374"),

    ("8.", "Open-Meteo. (2024). Historical weather API (ERA5) documentation. Retrieved from "
           "https://open-meteo.com/en/docs/historical-weather-api"),

    ("9.", "Abadi, M., Agarwal, A., Barham, P., Brevdo, E., Chen, Z., Citro, C., ... & Zheng, X. "
           "(2016). TensorFlow: A system for large-scale machine learning. In 12th USENIX "
           "Symposium on Operating Systems Design and Implementation (OSDI 16) (pp. 265-283)."),

    ("10.", "Streamlit Inc. (2024). Streamlit documentation. Retrieved from https://docs.streamlit.io"),

    ("11.", "Bradski, G. (2000). The OpenCV library. Dr. Dobb's Journal of Software Tools, 25, 120-125."),

    ("12.", "NIST. (2008). Secure hash standard (SHS). Federal Information Processing Standards "
            "Publication 180-3. National Institute of Standards and Technology. "
            "https://csrc.nist.gov/publications/detail/fips/180/3/final"),
]

for num, ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run(f"{num}  {ref}")
    run.font.size = Pt(11.5)
    run.font.name = 'Times New Roman'

add_page_break(doc)


# =============================================================================
# SAVE
# =============================================================================

out = "Project_Report_Kumar_Verma_MCA.docx"
doc.save(out)
print("Report saved: " + out)
print("Location: " + out)
print("Pages: approximately 60")
